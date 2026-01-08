import io
import pandas as pd
from docx import Document
from docx.shared import Pt
from pptx import Presentation
from reportlab.lib.pagesizes import letter, landscape
from reportlab.pdfgen import canvas
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
from reportlab.lib import colors
import json

def convert(file, target_type):
    """
    Convert CSV/XLS/XLSX files to various formats.
    Supports: CSV/XLS/XLSX → TXT, PDF, DOCX, JSON
    """
    try:
        output = io.BytesIO()
        
        # Try to read as Excel first, then fall back to CSV
        try:
            df = pd.read_excel(file)
        except:
            file.seek(0)  # Reset file pointer
            df = pd.read_csv(file)

        if target_type.upper() == 'TXT':
            # Convert to CSV format text
            output.write(df.to_csv(index=False).encode('utf-8'))
        
        elif target_type.upper() == 'PDF':
            # Create PDF with table
            pdf_output = io.BytesIO()
            doc = SimpleDocTemplate(pdf_output, pagesize=landscape(letter))
            
            # Prepare table data
            data = [df.columns.tolist()] + df.values.tolist()
            
            # Limit data if too large
            if len(data) > 50:
                data = data[:50]
            
            # Create table
            table = Table(data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
            ]))
            
            doc.build([table])
            pdf_output.seek(0)
            output.write(pdf_output.read())
        
        elif target_type.upper() == 'DOCX':
            # Convert to DOCX with table
            doc = Document()
            doc.add_heading('Data Table', level=1)
            
            # Add table
            table = doc.add_table(rows=len(df) + 1, cols=len(df.columns))
            table.style = 'Light Grid Accent 1'
            
            # Add headers
            for i, column in enumerate(df.columns):
                table.rows[0].cells[i].text = str(column)
            
            # Add data
            for i, row in enumerate(df.values):
                for j, value in enumerate(row):
                    table.rows[i + 1].cells[j].text = str(value)
            
            doc.save(output)
        
        elif target_type.upper() == 'JSON':
            # Convert to JSON
            json_data = df.to_dict(orient='records')
            output.write(json.dumps(json_data, indent=2).encode('utf-8'))
        
        elif target_type.upper() in ['XLS', 'XLSX']:
            # Convert to Excel
            df.to_excel(output, index=False, engine='openpyxl')
        
        elif target_type.upper() == 'CSV':
            # Convert to CSV
            output.write(df.to_csv(index=False).encode('utf-8'))
        
        else:
            raise ValueError(f"Unsupported target format: {target_type}")

        output.seek(0)
        return output
    
    except Exception as e:
        raise Exception(f"CSV/Excel conversion error: {str(e)}")
